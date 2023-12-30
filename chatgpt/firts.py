from docx import Document
import os

# File paths for the documents
doc_paths = [
    "/mnt/data/DISASTER MANAGEMENT.docx",
    "/mnt/data/FC _ BC.docx",
    "/mnt/data/Health and Hygiene.docx",
    "/mnt/data/MAP READING (2).docx",
    "/mnt/data/Military History.docx",
    "/mnt/data/NCC(GENERAL AWARENESS).docx",
    "/mnt/data/NCC LEADERSHIP AND PERSONALITY DEVELOPMENT.docx",
    "/mnt/data/NCC MCQ ARMED FORCES.docx",
    "/mnt/data/Wepons.docx"
]

# Function to extract questions from a document
def extract_questions(doc_path):
    doc = Document(doc_path)
    questions = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith('Q') or text.startswith('(') or text.endswith('?'):
            questions.append(text)
    return questions

# Extracting questions from each document
all_questions = []
for path in doc_paths:
    questions = extract_questions(path)
    all_questions.extend(questions)

# Selecting 50 questions for the question paper
selected_questions = all_questions[:50] if len(all_questions) >= 50 else all_questions

# Create a new document for the question paper
question_paper = Document()
question_paper.add_heading('Question Paper', 0)

# Adding questions to the document
for i, question in enumerate(selected_questions, 1):
    question_paper.add_paragraph(f'Q{i}: {question}')

# Adding a page break before the answer key
question_paper.add_page_break()

# Add heading for the answer key
question_paper.add_heading('Answer Key', 0)

# Placeholder for answers (to be filled in manually or with actual answers if available)
for i in range(1, len(selected_questions) + 1):
    question_paper.add_paragraph(f'Q{i}: [Answer]')

# Save the question paper to a file
question_paper_path = '/mnt/data/Question_Paper.docx'
question_paper.save(question_paper_path)

question_paper_path
